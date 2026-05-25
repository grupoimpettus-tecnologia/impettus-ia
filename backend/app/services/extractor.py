from pathlib import Path
import re
import pandas as pd
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook


def clean_cell(value) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    text = re.sub(r"\s+", " ", text)
    return text


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(file_path)
    if suffix == ".docx":
        return extract_docx(file_path)
    if suffix in [".xlsx", ".xlsm"]:
        return extract_excel_openpyxl(file_path)
    if suffix == ".xls":
        return extract_excel_pandas(file_path)
    if suffix == ".csv":
        return extract_csv(file_path)
    if suffix in [".txt", ".md"]:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Formato não suportado: {suffix}")


def extract_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    pages = []
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"\n--- Página {idx} ---\n{text}")
    return "\n".join(pages).strip()


def extract_docx(file_path: Path) -> str:
    doc = Document(str(file_path))
    blocks = []
    for p in doc.paragraphs:
        if p.text.strip():
            blocks.append(p.text.strip())
    for t_index, table in enumerate(doc.tables, start=1):
        blocks.append(f"\n--- Tabela DOCX {t_index} ---")
        for row in table.rows:
            cells = [clean_cell(cell.text) for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return "\n".join(blocks).strip()


def extract_excel_openpyxl(file_path: Path) -> str:
    """Extração melhorada para planilhas: preserva abas, células mescladas, fórmulas, linhas, colunas e pares Cabeçalho: Valor."""
    wb_values = load_workbook(file_path, data_only=True, read_only=False)
    wb_formulas = load_workbook(file_path, data_only=False, read_only=False)
    blocks = [f"Arquivo de planilha: {file_path.name}"]

    for ws in wb_values.worksheets:
        wsf = wb_formulas[ws.title]
        blocks.append(f"\n=== ABA: {ws.title} ===")
        max_row = min(ws.max_row or 0, 2000)
        max_col = min(ws.max_column or 0, 80)
        if max_row == 0 or max_col == 0:
            blocks.append("Aba vazia.")
            continue

        rows = []
        for r in range(1, max_row + 1):
            row_values = []
            for c in range(1, max_col + 1):
                v = clean_cell(ws.cell(r, c).value)
                f = clean_cell(wsf.cell(r, c).value)
                if f.startswith("=") and f != v:
                    v = f"{v} [fórmula: {f}]" if v else f"[fórmula: {f}]"
                row_values.append(v)
            if any(row_values):
                rows.append((r, row_values))

        if not rows:
            blocks.append("Aba sem dados legíveis.")
            continue

        # Detecta cabeçalho provável pela linha com mais células preenchidas nos primeiros 20 registros.
        sample = rows[:20]
        header_row_number, headers = max(sample, key=lambda item: sum(1 for v in item[1] if v))
        headers = [h if h else f"Coluna_{idx}" for idx, h in enumerate(headers, start=1)]
        blocks.append(f"Cabeçalhos prováveis na linha {header_row_number}: " + " | ".join([h for h in headers if h]))

        for row_number, values in rows:
            if row_number == header_row_number:
                continue
            pairs = []
            for header, value in zip(headers, values):
                if value:
                    pairs.append(f"{header}: {value}")
            if pairs:
                blocks.append(f"Linha {row_number} | " + " ; ".join(pairs))

    return "\n".join(blocks).strip()


def extract_excel_pandas(file_path: Path) -> str:
    xls = pd.ExcelFile(file_path)
    blocks = [f"Arquivo de planilha: {file_path.name}"]
    for sheet in xls.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet, dtype=str).fillna("")
        blocks.append(f"\n=== ABA: {sheet} ===")
        blocks.append("Colunas: " + " | ".join(map(str, df.columns)))
        for idx, row in df.iterrows():
            pairs = [f"{col}: {row[col]}" for col in df.columns if str(row[col]).strip()]
            if pairs:
                blocks.append(f"Linha {idx + 2} | " + " ; ".join(pairs))
    return "\n".join(blocks).strip()


def extract_csv(file_path: Path) -> str:
    try:
        df = pd.read_csv(file_path, dtype=str, sep=None, engine="python").fillna("")
    except Exception:
        df = pd.read_csv(file_path, dtype=str).fillna("")
    blocks = [f"Arquivo CSV: {file_path.name}", "Colunas: " + " | ".join(map(str, df.columns))]
    for idx, row in df.iterrows():
        pairs = [f"{col}: {row[col]}" for col in df.columns if str(row[col]).strip()]
        if pairs:
            blocks.append(f"Linha {idx + 2} | " + " ; ".join(pairs))
    return "\n".join(blocks).strip()
